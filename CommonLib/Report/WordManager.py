import CommonLib.Common as common
import CommonLib.Performance.PerformanceManager as pfm
import numpy as np

from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Cm


result_Img_width = 16
result_Img_height = 5
result_font_size = 15

timeSeries_Xlabel = 'Time'
timeSeries_Ylabel = 'Data'

title_Color = RGBColor(76, 89, 109)
title_size = Pt(24)
normal_font_Color = RGBColor(76, 89, 109)
normal_font_size = Pt(11)
small_font_size = Pt(8)

result_report_path = 'Data/Report/ML_Result/'   # 결과 워드문서
result_reportImg_path = 'Data/Report/ML_IMG/'   # 결과 Convert 이미지
result_image_path = 'Data/Image/ReportResult/'  # 분석결과 이미지

table_col_title = 'Report_title'
table_col_file = 'Report_File'
table_col_writedate = 'WriteDate'
table_col_userid = 'UserID'
table_col_imageYN = 'ImageYN'

# mongoDB 콜렉션
collection_Prophet = 'Prophet'
collection_Linear = 'LinearRegression'


class Cls_Paragraphs:
    """
    Word 텍스트 내에 변경할 문자내용 정의 Class
    파일명 : MLReport_Result_시퀀스.docx
    """
    def __init__(self, _text, _replaceText, _color, _fontSize, bold=False):
        self.text = _text
        self.replaceText = _replaceText
        self.color = _color
        self.fontSize = _fontSize
        self.bold = bold


def Create_Analysis_Report(titleVal, tableVal, targetDocumentFileName):
    """
    분석결과 리포트를 만든다
    :param targetDocumentFileName: 보고서 원본 대상 파일
    :param titleVal: 타이틀
    :param tableVal: 내용
    :return:
    """
    try:
        word_path = common.Path_Prj_Main() + result_report_path + targetDocumentFileName
        document = Document(word_path)
        fileName = Path(targetDocumentFileName).stem + '_' + common.Regexp_OnlyNumberbyDate() + '.docx'
        fileFullName = common.Path_Prj_Main() + result_report_path + fileName
        for item in titleVal:
            for txtPos in document.paragraphs:
                if item.text in txtPos.text:
                    common.Word_Make_Sentense(txtPos, item)
                    break

        for item in tableVal:
            for table in document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for tablePos in cell.paragraphs:
                            if item.text in tablePos.text:
                                common.Word_Make_Sentense(tablePos, item)

        document.save(fileFullName)
        return fileFullName
    except Exception as err:
        common.exception_print(err)


# def CreateReport_ARIMA(df, resultImagePath):
#     """
#     ARIMA 시계열 분석결과 보고서
#     :param df:
#     :param resultImagePath:
#     :return:
#     """
#     try:
#         print(df)
#         print(resultImagePath)
#         df_origin = pfm.descriptive_statistics(np.array(df['0'].dropna()))
#         df_yhat = pfm.descriptive_statistics(np.array(df['1'].dropna()))
#
#         titleVal = [Cls_Paragraphs("{타이틀}", "ARIMA 시계열 결과 보고서", title_Color, title_size, bold=True)]
#         tableVal = [
#             Cls_Paragraphs('{모델}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{분석방법}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{분석결과설명}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표1}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표2}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표3}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표4}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표설명1}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표설명2}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표설명3}', '', title_Color, normal_font_size),
#             Cls_Paragraphs('{지표설명4}', '', title_Color, normal_font_size)]
#
#         item_arr = ['평균', '중앙값', '누적합', '최대값']
#         for index, item in enumerate(item_arr):
#             Prophet_ParagraphsReturn(tableVal, index, item, df)
#
#         reportFileName = Create_Analysis_Report(titleVal, tableVal, 'MLReport_ARIMA.docx')
#         Result_Image_Add(reportFileName, resultImagePath)
#         return reportFileName
#
#     except Exception as err:
#         common.exception_print(err)


def TimeSeries_Pred(arr, df):
    for idx in range(1, 14):
        name = df.loc[idx-1, 'type']
        result = df.loc[idx-1, 'sts_result']
        arr.append(Cls_Paragraphs("{지표T" + str(idx) + "}", str(name), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{지표결과" + str(idx) + "}", str(result), title_Color, normal_font_size))


def TimeSeries_Result(arr, df):
    for idx in range(1, 32):
        date = df[['predic_date', 'predic_yhat']].tail(31).reset_index(drop=True).loc[idx - 1, 'predic_date']
        result = df[['predic_date', 'predic_yhat']].tail(31).reset_index(drop=True).loc[idx - 1, 'predic_yhat']
        arr.append(Cls_Paragraphs("{결과일자" + str(idx) + "}", str(date), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{결과값" + str(idx) + "}", str(result), title_Color, normal_font_size))


def CreateReport_Prophet(df, resultImagePath):
    """
    Prophet 시계열 분석결과 보고서
    참고블로그 : https://be-favorite.tistory.com/64
    :param df:
    :param resultImagePath:
    :return:
    """
    try:
        titleName = "Prophet 시계열 결과 보고서"
        Rsquared = df.at[0, 'sts_result']
        MAE = df.at[1, 'sts_result']
        MSE = df.at[2, 'sts_result']
        MSLE = df.at[4, 'sts_result']

        sts_df = pfm.descriptive_statistics(np.array(df['predic_yhat'].dropna()))

        df_yhat = df['predic_yhat'].dropna()
        df_yhat_date = df['predic_date'].dropna()
        df_yhat.reset_index(drop=True, inplace=True)
        df_yhat_date.reset_index(drop=True, inplace=True)

        titleVal = [Cls_Paragraphs("{타이틀}", titleName, title_Color, title_size, bold=True)]
        tableVal = [
            Cls_Paragraphs('{모델}', 'Facebook''s Prophet Model', title_Color, normal_font_size),
            Cls_Paragraphs('{분석방법}', 'Generalized additive models', title_Color, normal_font_size),
            Cls_Paragraphs('{라이브러리}', 'Prophet 패키지', title_Color, normal_font_size),
            Cls_Paragraphs('{일자}', common.ToDay(onlyDate=False), title_Color, normal_font_size),
            Cls_Paragraphs('{모델설명}', '과거에서부터 현재까지의 데이터를 바탕으로 미래에 대한 추세를 분석', title_Color, normal_font_size),
            Cls_Paragraphs('{지표1}', "R-Squared", title_Color, normal_font_size),
            Cls_Paragraphs('{지표2}', 'MAE', title_Color, normal_font_size),
            Cls_Paragraphs('{지표3}', 'MSE', title_Color, normal_font_size),
            Cls_Paragraphs('{지표4}', 'MSLE', title_Color, normal_font_size),
            Cls_Paragraphs('{지표설명1}', Rsquared, title_Color, normal_font_size),
            Cls_Paragraphs('{지표설명2}', MAE, title_Color, normal_font_size),
            Cls_Paragraphs('{지표설명3}', MSE, title_Color, normal_font_size),
            Cls_Paragraphs('{지표설명4}', MSLE, title_Color, normal_font_size)]

        result_arr = ['평균', '중앙값', '누적합', '최대값', '최소값', '범위', '분산', '표준편차', '1분위수(25%)',
                      '2분위수(25%)', '3분위수(25%)', '첨도', '왜도']

        for index, item in enumerate(result_arr):
            Prophet_ParagraphsReturn(tableVal, index, item, sts_df)

        for index in range(1, 32):
            Prophet_PredicData(tableVal, index, df_yhat_date, df_yhat)

        report_file_name = Create_Analysis_Report(titleVal, tableVal, 'MLReport_Prophet.docx')
        Result_Image_Add(report_file_name, resultImagePath)

        report_dictionary = {tableVal[i].text: tableVal[i].replaceText for i in range(len(tableVal))}
        report_dictionary[table_col_title] = titleName
        # report_dictionary[table_col_file] = report_file_name
        report_dictionary[table_col_file] = Path(report_file_name).stem
        report_dictionary[table_col_writedate] = common.ToDay(onlyDate=False)
        report_dictionary[table_col_userid] = 'seonghan.park@cwit.co.kr'
        report_dictionary[table_col_imageYN] = 'N'

        # 보고서 내용 저장
        common.Mongo_Insert_by_Dictionary(collection_Prophet, report_dictionary)

        return report_file_name

    except Exception as err:
        common.exception_print(err)


def Prophet_PredicData(tableArr, index, df_Date, df_y):
    try:
        tableArr.append(Cls_Paragraphs('{결과일자' + str(index) + '}', df_Date.at[index - 1], title_Color, normal_font_size))
        tableArr.append(Cls_Paragraphs('{결과값' + str(index) + '}', df_y.at[index - 1], title_Color, normal_font_size))
    except Exception as err:
        common.exception_print(err)


def Prophet_ParagraphsReturn(listArr, index, val, df):
    listArr.append(Cls_Paragraphs('{지표T' + str(index+1) + '}', val, title_Color, normal_font_size))
    listArr.append(
        Cls_Paragraphs('{지표결과' + str(index+1) + '}', df.at[index, 'sts_result'], title_Color, normal_font_size))


def CreateReport_LinearRegression(df, resultImagePath):
    """
    선형회귀 분석결과 보고서
    :param resultImagePath:
    :param df:
    :return:
    """
    try:
        titleName = "선형회귀 분석결과 보고서"
        Rsquared = df.at[0, 'P_MODEL_RESULT']
        MAE = df.at[1, 'P_MODEL_RESULT']
        MSE = df.at[2, 'P_MODEL_RESULT']
        MSLE = df.at[4, 'P_MODEL_RESULT']
        Summary = df.at[0, 'SUMMARY']

        titleVal = [Cls_Paragraphs("{타이틀}", titleName, title_Color, title_size, bold=True)]
        tableVal = [Cls_Paragraphs("{모델}", "Ordinary Least Squares Regression", title_Color, normal_font_size),
                    Cls_Paragraphs("{분석방법}", "최소자승법", title_Color, normal_font_size),
                    Cls_Paragraphs("{라이브러리}", "statsmodels 패키지", title_Color, normal_font_size),
                    Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
                    Cls_Paragraphs("{모델설명}", "오차를 최소화하여 회귀계수를 추정    ", title_Color, normal_font_size),
                    Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
                    Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
                    Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
                    Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
                    Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
                    Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
                    Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
                    Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
                    Cls_Paragraphs("{수치1}", common.front_Rexp(Summary, "No. Observations:", "AIC"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치2}", common.front_Rexp(Summary, "Df Residuals:", "BIC"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치3}", common.front_Rexp(Summary, "Df Model:", "Covariance Type"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치4}", common.front_Rexp(Summary, "R-squared:", "Model"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치5}", common.front_Rexp(Summary, "Adj. R-squared:", "Method"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치6}", common.front_Rexp(Summary, "F-statistic:", "Date"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치7}", common.front_Rexp(Summary, "Prob (F-statistic):", "Time"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치8}", common.front_Rexp(Summary, "AIC:", "Df Residuals"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치9}", common.front_Rexp(Summary, "BIC:", "Df Model"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치10}", common.front_Rexp(Summary, "Omnibus:", "Durbin-Watson"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치11}", common.front_Rexp(Summary, "Prob(Omnibus):", "Jarque-Bera (JB)"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치12}", common.front_Rexp(Summary, "Skew:", "Prob(JB)"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치13}", common.front_Rexp(Summary, "Kurtosis:", "Cond. No."), title_Color, small_font_size),
                    Cls_Paragraphs("{수치14}", common.front_Rexp(Summary, "Durbin-Watson:", "Prob(Omnibus)"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치15}", common.front_Rexp(Summary, "Jarque-Bera (JB):", "Skew"), title_Color, small_font_size),
                    Cls_Paragraphs("{수치16}", common.front_Rexp(Summary, "Cond. No.", "Notes:").replace('=', ''), title_Color, small_font_size)]

        report_file_name = Create_Analysis_Report(titleVal, tableVal, 'MLReport_Linear.docx')
        Result_Image_Add(report_file_name, resultImagePath)

        report_dictionary = {tableVal[i].text: tableVal[i].replaceText for i in range(len(tableVal))}
        report_dictionary[table_col_title] = titleName
        # report_dictionary[table_col_file] = report_file_name Full경로에서 파일명만 저장되도록 수정
        report_dictionary[table_col_file] = Path(report_file_name).stem
        report_dictionary[table_col_writedate] = common.ToDay(onlyDate=False)
        report_dictionary[table_col_userid] = 'seonghan.park@cwit.co.kr'
        report_dictionary[table_col_imageYN] = 'N'

        # 보고서 내용 저장
        common.Mongo_Insert_by_Dictionary(collection_Linear, report_dictionary)

        return report_file_name

    except Exception as err:
        common.exception_print(err)


def Result_Image_Add(targetDocumentName, resultfileName):
    """
    분석결과 이미지 추가
    :param targetDocumentName: 수정 타겟 워드 파일
    :param resultfileName: 문서에 추가 timeSeries_Ylabel가 할 결과이미지 파일
    :return:
    """
    try:
        if resultfileName == "":
            return

        # 결과 이미지 추가 파일
        doc = Document(targetDocumentName)
        tables = doc.tables

        imagePos = tables[2].rows[0].cells[1].add_paragraph()
        imagerun = imagePos.add_run()

        # 결과 이미지 추가
        imagerun.add_picture(resultfileName, width=Cm(result_Img_width), height=Cm(result_Img_height))
        doc.save(targetDocumentName)
    except Exception as err:
        common.exception_print(err)

# def CreateConvertImage(targetFileFullPath):
#     """
#     유료컴포넌트 사용 : aspose
#     분석결과 완료 된 docx 문서를 이미지로 변환
#     :param targetFileFullPath:분석완료 된 보고서의 full Path
#     :return:
#     """
#     import aspose.words as aw  # pip install aspose-words
#     # load document
#     doc = aw.Document(targetFileFullPath)
#     fileName = Path(targetFileFullPath).stem
#
#     # set output image format
#     options = aw.saving.ImageSaveOptions(aw.SaveFormat.PNG)
#
#     # loop through pages and convert them to PNG images
#     for pageNumber in range(doc.page_count):
#         options.page_set = aw.saving.PageSet(pageNumber)
#         doc.save(common.Path_Prj_Main() + result_reportImg_path + fileName + '_img.png', options)
#         # doc.save(common.Path_Prj_Main() + result_reportImg_path + fileName + '_img_' + str(pageNumber) + '.png', options)
#         break


if __name__ == '__main__':
    try:
        Rsquared = 123
        MAE = 325
        MSE = 4565
        MSLE = 1233

        titleVal = [Cls_Paragraphs("{타이틀}", "선형회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
        tableVal = [
            Cls_Paragraphs("{모델}", "Ordinary Least Squares Regression", title_Color, normal_font_size),
            Cls_Paragraphs("{분석방법}", "최소자승법", title_Color, normal_font_size),
            Cls_Paragraphs("{라이브러리}", "statsmodels 패키지", title_Color, normal_font_size),
            Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
            Cls_Paragraphs("{모델설명}", "오차를 최소화하여 회귀계수를 추정", title_Color, normal_font_size),
            Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
            Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
            Cls_Paragraphs("{분석결과설명}",
                           """
                           ①결과에 대한 분석전, Durbin watson을 확인한 결과 1.780으로 2에 근접하여
                           다중회귀분석모형에 적합하다고 판단하였다.
                           ②그리고 유의확률이 .000으로 경로 중
                           한가지 이상이 유효할 것이라것을 확인할 수 있었다.
                           ③ 또한 변수간 상관관계(R=.712)로 확인되었고
                           ④ 공차와 VIF 각각 0.1이상 10미만으로 다중공선성이 없는 것으로 확인되었다.
                           ⑤ 다음으로 각 경로에 유의성을 확인한 결과 신체자발성을 제외한
                           사회자발성(p<.001)과 인지자발성(p<.001)이 기질에 미치는 영향이 유효한 것으로 확인되었다.
                           ⑥ 유의한 변수에 대한 비표준화계수를 확인한 결과 사회자발셩(B=.283), 인지자발성(B=.369) 모두 양수로써
                           사회자발성과 인지자발성이 향상될수록 기질이 높아진 다는 것을 알 수 있었다.
                           ⑦또한, 사회자발성(베타=.375)과 인지자발성(베타=.445)이
                           기질에 미치는 영향력은 인지자발성, 사회자발성 순임을 알 수 있었다.
                           ⑧ 마지막으로 독립변수에 의해 종속변수가 설명되는 설명력은 50.7%임을 확인할 수 있었다.
                           """, title_Color, Pt(9))]

        word_path = 'D:/26.DataAnalisyPrj/ProxyServer/FastAPIServer/Data/Report/ML_Result/MLReport_Linear.docx'
        document = Document(word_path)
        fileName = 'MLReport_Linear_' + common.Regexp_OnlyNumberbyDate() + '.docx'
        fileFullName = common.Path_Prj_Main() + result_report_path + fileName

        for item in titleVal:
            for txtPos in document.paragraphs:
                if item.text in txtPos.text:
                    common.Word_Make_Sentense(txtPos, item)
                    break

        for item in tableVal:
            for table in document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for tablePos in cell.paragraphs:
                            if item.text in tablePos.text:
                                common.Word_Make_Sentense(tablePos, item)

        document.save(fileFullName)
    except Exception as err:
        common.exception_print(err)